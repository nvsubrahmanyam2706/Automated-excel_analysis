# report_generator.py
from docx import Document
from docx.shared import Inches
from io import BytesIO
import tempfile
import os
 
def generate_report(df, prompt, results_text, visualizations):
    """Generate a complete Word report including visuals & analysis."""
    doc = Document()
    doc.add_heading("📊 AI Excel Analysis Report", 0)
 
    # --- Prompt & Dataset Info ---
    doc.add_heading("User Prompt:", level=1)
    doc.add_paragraph(prompt)
 
    doc.add_heading("Dataset Info:", level=1)
    doc.add_paragraph(f"Shape: {df.shape}")
    doc.add_paragraph("Missing Values:")
    doc.add_paragraph(df.isnull().sum().to_string())
 
    doc.add_heading("Dataset Preview (first 5 rows):", level=1)
    doc.add_paragraph(df.head().to_string())
 
    # --- Gemini Analysis Results ---
    doc.add_heading("Gemini Analysis Results:", level=1)
    doc.add_paragraph(results_text if results_text.strip() else "No results captured.")
 
    # --- Visualizations ---
    if visualizations:
        doc.add_heading("Visualizations:", level=1)
        for i, vis in enumerate(visualizations):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                tmp_file.write(vis)
                tmp_path = tmp_file.name
            doc.add_picture(tmp_path, width=Inches(5.5))
            os.remove(tmp_path)
    else:
        doc.add_paragraph("No visualizations available.")
 
    # --- Finalize report ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
 