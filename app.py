# app.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from prophet import Prophet
import xgboost as xgb
import statsmodels.api as sm
from docx import Document
from io import BytesIO
from llm_helper import configure_gemini, suggest_algorithms, generate_code_for_algorithm
 
# --- Streamlit Setup ---
st.set_page_config(page_title="AI Excel Analyzer (Gemini)", layout="wide")
st.title("🤖 Gemini-Driven Excel Data Analysis")
 
# --- Gemini API Key ---
api_key = st.text_input("🔑 Enter Gemini API Key:", type="password")
if not api_key:
    st.stop()
configure_gemini(api_key)
 
# --- State Variables ---
if "results_text" not in st.session_state:
    st.session_state.results_text = ""
if "visualizations" not in st.session_state:
    st.session_state.visualizations = []
 
# --- Step 1: Upload Excel ---
uploaded_file = st.file_uploader("📂 Upload Excel file", type=["xls", "xlsx"])
if uploaded_file:
    df = pd.read_excel(uploaded_file)
 
    # Clean and reset index for safety
    df = df.reset_index(drop=True)
 
    st.write("### Dataset Preview")
    st.dataframe(df.head())
    st.write("**Shape:**", df.shape)
    st.write("**Missing values:**")
    st.write(df.isnull().sum())
 
    # --- Step 2: User Prompt ---
    prompt = st.text_input("💬 Enter your analysis prompt")
    if prompt:
        st.info("Generating algorithm suggestions from Gemini...")
        algos = suggest_algorithms(prompt, df.head())
        st.write("### 🤖 Suggested Algorithms")
 
        cols = st.columns(len(algos))
        for i, algo in enumerate(algos):
            if cols[i].button(algo):
                st.info(f"Generating and executing {algo} code via Gemini...")
 
                code = generate_code_for_algorithm(prompt, df.head(), algo)
 
                # Capture visuals
                captured_plots = []
 
                def capture_pyplot(fig=None):
                    """Capture st.pyplot() calls for report."""
                    buf = BytesIO()
                    if fig is None:
                        fig = plt.gcf()
                    fig.savefig(buf, format="png", bbox_inches="tight")
                    buf.seek(0)
                    captured_plots.append(buf.read())
                    st.image(buf)
                    plt.close(fig)
 
                # --- Clean Data before use ---
                df_cleaned = df.copy()
                for col in df_cleaned.select_dtypes(include=np.number).columns:
                    df_cleaned[col].fillna(df_cleaned[col].median(), inplace=True)
                for col in df_cleaned.select_dtypes(exclude=np.number).columns:
                    df_cleaned[col].fillna(df_cleaned[col].mode()[0], inplace=True)
                df_cleaned.dropna(how="all", inplace=True)
                df_cleaned.reset_index(drop=True, inplace=True)
 
                # --- Create local environment for Gemini ---
                local_env = {
                    "df": df_cleaned.copy(),
                    "pd": pd,
                    "st": st,
                    "plt": plt,
                    "sns": sns,
                    "np": np,
                    "train_test_split": train_test_split,
                    "LabelEncoder": LabelEncoder,
                    "OneHotEncoder": OneHotEncoder,
                    "ColumnTransformer": ColumnTransformer,
                    "Pipeline": Pipeline,
                    "DecisionTreeClassifier": DecisionTreeClassifier,
                    "RandomForestClassifier": RandomForestClassifier,
                    "KMeans": KMeans,
                    "DBSCAN": DBSCAN,
                    "AgglomerativeClustering": AgglomerativeClustering,
                    "Prophet": Prophet,
                    "sm": sm,
                    "xgb": xgb,
                }
 
                local_env["st"].pyplot = capture_pyplot
 
                # --- Run Gemini Code ---
                try:
                    exec(code, local_env)
                    st.success(f"{algo} executed successfully! ✅")
 
                    st.session_state.results_text += f"\n\n### {algo} Analysis\nGemini successfully executed {algo}.\n"
                    st.session_state.visualizations.extend(captured_plots)
 
                except ValueError as ve:
                    st.error(f"ValueError: {ve}")
                    if "Unalignable boolean Series" in str(ve):
                        st.warning("⚠️ Gemini tried to filter using misaligned boolean indexes. We’ll fix index alignment next time.")
                except Exception as e:
                    st.error(f"Unexpected error running Gemini code: {e}")
 
        # --- Step 3: Generate Report ---
        st.write("### 📄 Generate & Download Report")
        if st.button("📥 Generate Word Report"):
            if not st.session_state.results_text and not st.session_state.visualizations:
                st.warning("⚠️ Run at least one algorithm before generating the report.")
            else:
                doc = Document()
                doc.add_heading("AI Excel Analysis Report", 0)
                doc.add_paragraph(f"User Prompt: {prompt}")
                doc.add_paragraph(f"Dataset Shape: {df.shape}")
                doc.add_heading("Missing Values:", level=1)
                doc.add_paragraph(df.isnull().sum().to_string())
                doc.add_heading("Dataset Preview (first 5 rows):", level=1)
                doc.add_paragraph(df.head().to_string())
                doc.add_heading("Gemini Insights:", level=1)
                doc.add_paragraph(st.session_state.results_text)
                doc.add_heading("Visualizations:", level=1)
                for img_bytes in st.session_state.visualizations:
                    img_stream = BytesIO(img_bytes)
                    doc.add_picture(img_stream)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="📄 Download Word Report",
                    data=buffer,
                    file_name="AI_Excel_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )